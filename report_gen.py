import os
import subprocess
import mysql.connector
from docx import Document
from datetime import datetime
from collections import Counter

# Obtener la ruta donde está este script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Crear carpeta 'Informes' en esa misma ruta
output_dir = os.path.join(script_dir, "Informes")
os.makedirs(output_dir, exist_ok=True)

# Mostrar la ruta para confirmar
print(f"Los informes se guardarán en: {output_dir}")

# Conexión a la base de datos
conn = mysql.connector.connect(
    host="localhost",
    user="root",
    password="root",
    database="honey"
)
cursor = conn.cursor(dictionary=True)
cursor.execute("SELECT * FROM logs_ataques")
registros = cursor.fetchall()
cursor.close()
conn.close()

# Procesar estadísticas
total_eventos = len(registros)
tipos_evento = Counter(r['event_type'] for r in registros)
ips_origen = Counter(r['ip_origen'] for r in registros)

# Crear documento
doc = Document()
doc.add_heading('Informe de Actividad del Honeypot', 0)

doc.add_paragraph("Este documento recoge la actividad registrada por el sistema honeypot conectado a Suricata. Se analizan los eventos detectados, las IPs más activas y los tipos de ataque observados.\n")
doc.add_paragraph("Fecha de generación: " + datetime.now().strftime("%d/%m/%Y %H:%M:%S"))

# Sección: Resumen general
doc.add_heading('Resumen General', level=1)
doc.add_paragraph(f"Total de eventos registrados: {total_eventos}")
doc.add_paragraph(f"IPs de origen únicas: {len(ips_origen)}")
doc.add_paragraph("Tipos de eventos detectados:")
for tipo, cantidad in tipos_evento.items():
    doc.add_paragraph(f"  • {tipo}: {cantidad} eventos", style='List Bullet')

# Sección: Análisis por IP de origen
doc.add_heading('Análisis por IP de Origen', level=1)
for ip, cantidad in ips_origen.most_common():
    eventos_ip = [r for r in registros if r['ip_origen'] == ip]
    tipos = Counter(r['event_type'] for r in eventos_ip)
    doc.add_heading(f"IP: {ip}", level=2)
    doc.add_paragraph(f"Total de eventos: {cantidad}")
    doc.add_paragraph("Tipos de eventos:")
    for tipo, c in tipos.items():
        doc.add_paragraph(f"  • {tipo}: {c}", style='List Bullet')
    doc.add_paragraph(f"Interpretación: La IP {ip} ha generado múltiples eventos, lo que puede indicar una actividad automatizada como escaneo de puertos o pruebas de vulnerabilidades.\n")

# Sección: Eventos clasificados por tipo
doc.add_heading('Eventos por Tipo', level=1)
for tipo, cantidad in tipos_evento.items():
    doc.add_heading(tipo.upper(), level=2)
    doc.add_paragraph(f"Se han registrado {cantidad} eventos de tipo {tipo}.")
    if tipo.lower() == 'dns':
        doc.add_paragraph("Los eventos DNS pueden indicar intentos de reconocimiento de red o exfiltración de datos mediante dominios controlados por el atacante.")
    elif tipo.lower() == 'http':
        doc.add_paragraph("Los eventos HTTP suelen estar relacionados con escaneos web, intentos de acceso a recursos inseguros o ejecución de comandos remotos.")
    else:
        doc.add_paragraph("Este tipo de evento debe ser investigado según su naturaleza específica y frecuencia.")

# Sección: Conclusiones
doc.add_heading('Conclusiones y Recomendaciones', level=1)
doc.add_paragraph("Los datos obtenidos muestran actividad sospechosa en la red del honeypot. Se recomienda:")
doc.add_paragraph("  • Analizar las IPs con mayor número de eventos y considerarlas para listas de bloqueo.")
doc.add_paragraph("  • Monitorizar continuamente la red para detectar patrones de comportamiento repetitivo.")
doc.add_paragraph("  • Reforzar medidas de seguridad como firewalls o sistemas IDS adicionales.")

# Guardar el DOCX
nombre_base = datetime.now().strftime("informe_%Y-%m-%d_%H-%M-%S")
docx_path = os.path.join(output_dir, f"{nombre_base}.docx")
doc.save(docx_path)
print(f"[✓] Informe DOCX creado: {docx_path}")

# Convertir a PDF con LibreOffice
print("[+] Convirtiendo a PDF...")
try:
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path], check=True)
    print(f"[✓] Informe PDF creado: {os.path.join(output_dir, f'{nombre_base}.pdf')}")
except FileNotFoundError:
    print("[!] LibreOffice no está instalado. Ejecuta: sudo apt install libreoffice")
